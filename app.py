import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
import re
import json
import shutil
import subprocess
import sys
from datetime import datetime, date
from pathlib import Path

# ── third-party (pip install python-docx openpyxl) ──────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    import openpyxl
except ImportError:
    messagebox.showerror(
        "Missing dependencies",
        "Run:  pip install python-docx openpyxl\nthen relaunch the app."
    )
    sys.exit(1)

# ── colours ──────────────────────────────────────────────────────────────────
BG        = "#F7F6F3"
SURFACE   = "#FFFFFF"
BORDER    = "#E2E0D8"
TEXT      = "#1A1918"
MUTED     = "#6B6963"
PURPLE    = "#3C3489"
PURPLE_LT = "#EEEDFE"
AMBER     = "#FAEEDA"
AMBER_BD  = "#FAC775"
AMBER_TX  = "#633806"
GREEN_LT  = "#E1F5EE"
GREEN_BD  = "#9FE1CB"
GREEN_TX  = "#085041"
RED_LT    = "#FCEBEB"
RED_BD    = "#F09595"
RED_TX    = "#791F1F"
GREEN_BTN = "#DDEFE8"
GREEN_BTN_ACTIVE = "#CBE5DB"
RED_BTN   = "#F5E2DF"
RED_BTN_ACTIVE = "#EDD0CC"
FONT      = ("Segoe UI", 10)
FONT_SM   = ("Segoe UI", 9)
FONT_BOLD = ("Segoe UI", 10, "bold")
FONT_H    = ("Segoe UI", 12, "bold")

PLACEHOLDER_RE = re.compile(r"\{\{(.+?)\}\}")
BUILTIN_TEMPLATE_FIELDS = ["Issuance Date"]
CONTRACT_OUTPUT_FOLDER = "Contract Generated"
NAME_FIELD_KEYS = {"full name", "candidate name", "name"}
BOLD_FIELD_KEYS = NAME_FIELD_KEYS | {"title", "job title", "joining date", "start date", "salary"}
SALARY_FIELD_KEYS = {"salary"}
DATE_ONLY_FIELD_KEYS = {"joining date", "start date"}

# ─────────────────────────────────────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────────────────────────────────────

class SolidActionButton(tk.Frame):
    """Flat button with exact colors, avoiding macOS Tk tinting."""

    def __init__(self, master, text, command, bg, fg, activebackground=None,
                 activeforeground=None, disabledforeground=None,
                 disabledbackground=None, highlightbackground=None,
                 font=None, padx=10, pady=4, cursor="hand2"):
        super().__init__(master, bg=highlightbackground or bg, bd=0, highlightthickness=0)
        self._command = command
        self._state = "normal"
        self._cursor = cursor
        self._hover = False
        self._pressed = False
        self._bg = bg
        self._fg = fg
        self._activebackground = activebackground or bg
        self._activeforeground = activeforeground or fg
        self._disabledforeground = disabledforeground or fg
        self._disabledbackground = disabledbackground or bg
        self._border = highlightbackground or bg

        self._label = tk.Label(self, text=text, font=font, bg=bg, fg=fg,
                               bd=0, padx=padx, pady=pady, cursor=cursor)
        self._label.pack(fill="both", expand=True, padx=1, pady=1)

        for widget in (self, self._label):
            widget.bind("<Enter>", self._on_enter)
            widget.bind("<Leave>", self._on_leave)
            widget.bind("<ButtonPress-1>", self._on_press)
            widget.bind("<ButtonRelease-1>", self._on_release)

        self._apply_style()

    def configure(self, cnf=None, **kwargs):
        if cnf:
            kwargs = {**cnf, **kwargs}
        if not kwargs:
            return super().configure()

        bg_changed = "bg" in kwargs
        border_changed = "highlightbackground" in kwargs

        if "command" in kwargs:
            self._command = kwargs.pop("command")
        if "state" in kwargs:
            self._state = kwargs.pop("state")
        if "bg" in kwargs:
            self._bg = kwargs.pop("bg")
        if "fg" in kwargs:
            self._fg = kwargs.pop("fg")
        if "activebackground" in kwargs:
            self._activebackground = kwargs.pop("activebackground")
        if "activeforeground" in kwargs:
            self._activeforeground = kwargs.pop("activeforeground")
        if "disabledforeground" in kwargs:
            self._disabledforeground = kwargs.pop("disabledforeground")
        if "disabledbackground" in kwargs:
            self._disabledbackground = kwargs.pop("disabledbackground")
        if "highlightbackground" in kwargs:
            self._border = kwargs.pop("highlightbackground")
        elif bg_changed and not border_changed:
            self._border = self._bg

        if "cursor" in kwargs:
            self._cursor = kwargs.pop("cursor")

        label_kwargs = {}
        for key in ("text", "font", "padx", "pady"):
            if key in kwargs:
                label_kwargs[key] = kwargs.pop(key)
        if label_kwargs:
            self._label.configure(**label_kwargs)

        self._apply_style()

    config = configure

    def _apply_style(self):
        if self._state == "disabled":
            border = self._disabledbackground
            fill = self._disabledbackground
            text = self._disabledforeground
            cursor = "arrow"
        else:
            if self._pressed or self._hover:
                border = self._activebackground
                fill = self._activebackground
                text = self._activeforeground
            else:
                border = self._border
                fill = self._bg
                text = self._fg
            cursor = self._cursor

        super().configure(bg=border, cursor=cursor)
        self._label.configure(bg=fill, fg=text, cursor=cursor)

    def _on_enter(self, _event):
        if self._state != "disabled":
            self._hover = True
            self._apply_style()

    def _on_leave(self, _event):
        if self._state != "disabled":
            self._hover = False
            self._pressed = False
            self._apply_style()

    def _on_press(self, _event):
        if self._state != "disabled":
            self._pressed = True
            self._apply_style()

    def _on_release(self, event):
        if self._state == "disabled":
            return
        inside = self.winfo_containing(event.x_root, event.y_root) in (self, self._label)
        self._pressed = False
        self._hover = inside
        self._apply_style()
        if inside and self._command:
            self._command()

def _sheet_headers(sheet):
    if sheet is None:
        return []
    first_row = next(sheet.iter_rows(min_row=1, max_row=1), None)
    if first_row is None:
        return []

    return [str(cell.value).strip() if cell.value else "" for cell in first_row]


def build_template_fields(extra_fields=None):
    ordered = []
    seen = set()
    for field in [*BUILTIN_TEMPLATE_FIELDS, *(extra_fields or [])]:
        name = str(field or "").strip()
        if not name or name in seen:
            continue
        seen.add(name)
        ordered.append(name)
    return ordered


def current_issuance_date():
    now = datetime.now()
    return f"{now.day} {now.strftime('%B %Y')}"


def resolve_contract_output_dir(path):
    raw = str(path or "").strip()
    if not raw:
        return ""
    if os.path.basename(os.path.normpath(raw)).casefold() == CONTRACT_OUTPUT_FOLDER.casefold():
        return raw
    return os.path.join(raw, CONTRACT_OUTPUT_FOLDER)


def is_candidate_name_field(name):
    return str(name or "").strip().casefold() in NAME_FIELD_KEYS


def is_bold_insert_field(name):
    return str(name or "").strip().casefold() in BOLD_FIELD_KEYS


def is_salary_field(name):
    return str(name or "").strip().casefold() in SALARY_FIELD_KEYS


def is_date_only_field(name):
    return str(name or "").strip().casefold() in DATE_ONLY_FIELD_KEYS


def _format_excel_value(header, value):
    if value is None:
        return ""

    if is_date_only_field(header):
        if isinstance(value, datetime):
            value = value.date()
        if isinstance(value, date):
            return f"{value.day} {value.strftime('%B %Y')}"

        text = str(value).strip()
        if not text:
            return ""
        text = re.sub(r"(?:[T\s]+00:00(?::00(?:\.0+)?)?)$", "", text)
        return text

    return str(value).strip()

def _find_candidate_sheet(workbook):
    for sheet in workbook.worksheets:
        name = sheet.title.lower().strip()
        if name == "annex b" or "annex b" in name or "candidate" in name:
            return sheet
    return workbook.worksheets[0] if workbook.worksheets else None


def _find_signatory_sheet(workbook):
    for sheet in workbook.worksheets:
        name = sheet.title.lower().strip()
        if name == "annex c" or "annex c" in name or "signatory" in name or "sign" in name:
            return sheet
    return workbook.worksheets[1] if len(workbook.worksheets) > 1 else None


def extract_excel_fields(path):
    """Return unique Excel column headers from the candidate and signatory sheets."""
    wb = openpyxl.load_workbook(path, read_only=True)
    try:
        ordered = []
        seen = set()
        for header in _sheet_headers(_find_candidate_sheet(wb)) + _sheet_headers(_find_signatory_sheet(wb)):
            if not header:
                continue
            if header not in seen:
                seen.add(header)
                ordered.append(header)
        return build_template_fields(ordered)
    finally:
        wb.close()


def find_default_template(base_dir=None):
    """Prefer a local .docx with 'template' in the filename, else the only .docx file."""
    root = Path(base_dir or Path(__file__).resolve().parent)
    docx_files = sorted(root.glob("*.docx"))
    if not docx_files:
        return ""

    preferred = [path for path in docx_files if "template" in path.stem.lower()]
    if preferred:
        return str(preferred[0])
    if len(docx_files) == 1:
        return str(docx_files[0])
    return ""


def normalize_placeholder_spacing(text):
    """Render placeholders with inner spacing for a more readable editor view."""
    return PLACEHOLDER_RE.sub(lambda m: f"{{{{ {m.group(1).strip()} }}}}", text)


def normalize_rank(value):
    """Normalize rank text so Annex B and Annex C can match reliably."""
    return re.sub(r"\s+", " ", str(value or "").strip()).casefold()


def _set_paragraph_text(paragraph, text):
    if paragraph.text == text:
        return
    if paragraph.runs:
        for i, run in enumerate(paragraph.runs):
            run.text = text if i == 0 else ""
    else:
        paragraph.add_run(text)


def _copy_run_format(source_run, target_run):
    target_run.style = source_run.style
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline

    target_font = target_run.font
    source_font = source_run.font
    target_font.name = source_font.name
    target_font.size = source_font.size
    target_font.bold = source_font.bold
    target_font.italic = source_font.italic
    target_font.underline = source_font.underline


def _apply_insert_field_format(run, field_name):
    make_bold = is_bold_insert_field(field_name)
    run.style = None
    run.bold = make_bold
    run.italic = False
    run.underline = False
    run.font.name = "Aptos"
    run.font.size = Pt(12)
    run.font.bold = make_bold
    run.font.italic = False
    run.font.underline = False


def _clear_paragraph_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)


def _slice_plain_segments(full_text, start, end, run_spans):
    if start >= end:
        return []
    if not run_spans:
        return [(full_text[start:end], None)]

    pieces = []
    for run, run_start, run_end in run_spans:
        seg_start = max(start, run_start)
        seg_end = min(end, run_end)
        if seg_start < seg_end:
            pieces.append((full_text[seg_start:seg_end], run))
    return pieces or [(full_text[start:end], None)]


def read_docx_plain_text(docx_path):
    doc = Document(docx_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def save_docx_plain_text(docx_path, text):
    """Write edited plain-text lines back into an existing docx while preserving untouched paragraphs."""
    doc = Document(docx_path)
    lines = text.split("\n")
    paragraphs = list(doc.paragraphs)

    for idx, line in enumerate(lines):
        if idx < len(paragraphs):
            _set_paragraph_text(paragraphs[idx], line)
        else:
            doc.add_paragraph(line)

    for paragraph in paragraphs[len(lines):]:
        if paragraph.text:
            _set_paragraph_text(paragraph, "")

    doc.save(docx_path)


def open_path(path):
    """Open a file or folder with the OS default application."""
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(path)])
        else:
            subprocess.Popen(["xdg-open", str(path)])
        return True, ""
    except Exception as exc:
        return False, str(exc)


def read_excel(path):
    """Return (candidates: list[dict], signatories: dict[rank->{'name','title'}])."""
    wb = openpyxl.load_workbook(path)

    # Annex B — candidate data
    b_sheet = _find_candidate_sheet(wb)

    headers = _sheet_headers(b_sheet)
    candidates = []
    for row in b_sheet.iter_rows(min_row=2, values_only=True):
        if all(v is None for v in row):
            continue
        candidates.append({headers[i]: _format_excel_value(headers[i], v) for i, v in enumerate(row)})

    # Annex C — signatory mapping
    c_sheet = _find_signatory_sheet(wb)

    signatories = {}
    if c_sheet:
        c_hdrs = _sheet_headers(c_sheet)
        for row in c_sheet.iter_rows(min_row=2, values_only=True):
            if all(v is None for v in row):
                continue
            rd = {c_hdrs[i]: _format_excel_value(c_hdrs[i], v) for i, v in enumerate(row)}
            raw_rank = rd.get("Rank") or rd.get("rank") or list(rd.values())[0]
            rank_key = normalize_rank(raw_rank)
            if not rank_key:
                continue
            if rank_key in signatories:
                raise ValueError(f"Duplicate Rank found in Annex C: {raw_rank}")
            signatories[rank_key] = rd

    return candidates, signatories


def extract_placeholders(docx_path):
    """Return sorted list of unique placeholder names found in the template."""
    doc = Document(docx_path)
    found = set()
    for para in doc.paragraphs:
        for m in PLACEHOLDER_RE.finditer(para.text):
            found.add(m.group(1).strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for m in PLACEHOLDER_RE.finditer(para.text):
                        found.add(m.group(1).strip())
    return sorted(found)


def fill_template(docx_path, replacements, out_path):
    """Deep-replace all {{key}} occurrences and save to out_path."""
    doc = Document(docx_path)

    def replace_in_para(para):
        full = para.text
        if "{{" not in full:
            return
        matches = list(PLACEHOLDER_RE.finditer(full))
        if not matches:
            return
        run_spans = []
        cursor = 0
        for run in para.runs:
            start = cursor
            end = start + len(run.text)
            run_spans.append((run, start, end))
            cursor = end

        pieces = []
        cursor = 0
        for match in matches:
            start, end = match.span()
            if cursor < start:
                for text, source_run in _slice_plain_segments(full, cursor, start, run_spans):
                    if text:
                        pieces.append(("plain", text, source_run, None))

            field_name = match.group(1).strip()
            replacement = str(replacements.get(field_name, match.group(0)))
            if replacement and is_salary_field(field_name) and pieces:
                prev_kind, prev_text, prev_source_run, prev_field_name = pieces[-1]
                if prev_kind == "plain" and prev_text.endswith("S$"):
                    prev_text = prev_text[:-2]
                    if prev_text:
                        pieces[-1] = (prev_kind, prev_text, prev_source_run, prev_field_name)
                    else:
                        pieces.pop()
                    replacement = f"S${replacement}"
            if replacement:
                pieces.append(("field", replacement, None, field_name))
            cursor = end

        if cursor < len(full):
            for text, source_run in _slice_plain_segments(full, cursor, len(full), run_spans):
                if text:
                    pieces.append(("plain", text, source_run, None))

        _clear_paragraph_runs(para)
        for kind, text, source_run, field_name in pieces:
            run = para.add_run(text)
            if kind == "plain" and source_run is not None:
                _copy_run_format(source_run, run)
            elif kind == "field":
                _apply_insert_field_format(run, field_name)

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    doc.save(out_path)


def docx_to_pdf(docx_path, out_dir):
    """Convert docx → pdf using LibreOffice (must be installed)."""
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        capture_output=True, text=True, timeout=60
    )
    stem = Path(docx_path).stem
    pdf = Path(out_dir) / f"{stem}.pdf"
    if pdf.exists():
        return str(pdf)
    raise RuntimeError(result.stderr or "PDF conversion failed")


# ─────────────────────────────────────────────────────────────────────────────
#  Custom widgets
# ─────────────────────────────────────────────────────────────────────────────

class FilePicker(tk.Frame):
    def __init__(self, parent, label, mode="file", filetypes=None, on_change=None, **kw):
        super().__init__(parent, bg=BG, **kw)
        self.mode = mode
        self.filetypes = filetypes or [("All files", "*.*")]
        self.on_change = on_change
        self.var = tk.StringVar()
        if self.on_change:
            self.var.trace_add("write", self._emit_change)

        tk.Label(self, text=label, font=FONT_SM, bg=BG, fg=MUTED).pack(anchor="w", pady=(0, 3))
        row = tk.Frame(self, bg=BG)
        row.pack(fill="x")

        entry = tk.Entry(row, textvariable=self.var, font=FONT, bg=SURFACE, fg=TEXT,
                         relief="flat", bd=0, highlightthickness=1,
                         highlightbackground=BORDER, highlightcolor=PURPLE)
        entry.pack(side="left", fill="x", expand=True, ipady=5, ipadx=6)

        btn = tk.Button(row, text="Browse", font=FONT_SM, bg=SURFACE, fg=TEXT,
                        relief="flat", bd=0, highlightthickness=1,
                        highlightbackground=BORDER, cursor="hand2",
                        padx=10, pady=4, command=self._browse)
        btn.pack(side="left", padx=(6, 0))

    def _emit_change(self, *_args):
        if self.on_change:
            self.on_change(self.var.get())

    def _browse(self):
        if self.mode == "file":
            path = filedialog.askopenfilename(filetypes=self.filetypes)
        else:
            path = filedialog.askdirectory()
        if path:
            self.var.set(path)

    def get(self):
        return self.var.get()

    def set(self, v):
        self.var.set(v)


class StatusDot(tk.Canvas):
    COLOURS = {
        "pending":       (BORDER, BORDER),
        "approved":      ("#1D9E75", "#1D9E75"),
        "needs_changes": ("#E24B4A", "#E24B4A"),
        "rejected":      ("#E24B4A", "#E24B4A"),
        "edited":        ("#BA7517", "#BA7517"),
    }

    def __init__(self, parent, status="pending", **kw):
        super().__init__(parent, width=10, height=10, bg=parent["bg"],
                         highlightthickness=0, **kw)
        self._status = status
        self._draw()

    def _draw(self):
        self.delete("all")
        fill, outline = self.COLOURS.get(self._status, (BORDER, BORDER))
        self.create_oval(1, 1, 9, 9, fill=fill, outline=outline, width=1)

    def set(self, status):
        self._status = status
        self._draw()

    def get(self):
        return self._status


# ─────────────────────────────────────────────────────────────────────────────
#  Generate Tab
# ─────────────────────────────────────────────────────────────────────────────

class GenerateTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._build()

    def _build(self):
        pad = dict(padx=20, pady=8)

        tk.Label(self, text="Contract Generator", font=FONT_H, bg=BG, fg=TEXT
                 ).pack(anchor="w", padx=20, pady=(10, 2))
        tk.Label(self, text="Select your data file and base output folder, then generate. Contracts are saved into 'Contract Generated'. Edit the template only if you need to adjust the wording.",
                 font=FONT_SM, bg=BG, fg=MUTED).pack(anchor="w", padx=20, pady=(0, 10))

        self.excel_picker = FilePicker(self, "Excel data file (Annex B & C)",
                                       filetypes=[("Excel files", "*.xlsx *.xlsm")],
                                       on_change=self._excel_changed)
        self.excel_picker.pack(fill="x", **pad)

        self.template_picker = FilePicker(self, "Word contract template",
                                          filetypes=[("Word files", "*.docx")],
                                          on_change=self._template_changed)
        self.template_picker.pack(fill="x", **pad)

        template_hint = tk.Frame(self, bg=BG)
        template_hint.pack(fill="x", padx=20, pady=(0, 6))
        tk.Label(template_hint, text="Optional:", font=FONT_SM, bg=BG, fg=MUTED
                 ).pack(side="left")
        tk.Button(template_hint, text="Edit Template", font=FONT_SM, bg=SURFACE, fg=TEXT,
                  relief="flat", bd=0, highlightthickness=1,
                  highlightbackground=BORDER, activebackground=SURFACE,
                  activeforeground=TEXT, padx=10, pady=4, cursor="hand2",
                  command=self._edit_template).pack(side="left", padx=(8, 0))

        self.output_picker = FilePicker(self, "Output base folder", mode="dir")
        self.output_picker.pack(fill="x", **pad)

        # Divider
        tk.Frame(self, bg=BORDER, height=1).pack(fill="x", padx=20, pady=8)

        # Progress
        tk.Label(self, text="Progress", font=FONT_SM, bg=BG, fg=MUTED
                 ).pack(anchor="w", padx=20)
        self.progress = ttk.Progressbar(self, mode="determinate", maximum=100)
        self.progress.pack(fill="x", padx=20, pady=(4, 2))
        self.progress_label = tk.Label(self, text="Ready", font=FONT_SM, bg=BG, fg=MUTED)
        self.progress_label.pack(anchor="w", padx=20)

        # Log
        tk.Label(self, text="Log", font=FONT_SM, bg=BG, fg=MUTED
                 ).pack(anchor="w", padx=20, pady=(10, 2))
        log_frame = tk.Frame(self, bg=SURFACE, highlightthickness=1,
                             highlightbackground=BORDER)
        log_frame.pack(fill="x", padx=20)
        self.log = tk.Text(log_frame, height=6, font=("Consolas", 9),
                           bg=SURFACE, fg=MUTED, relief="flat", bd=4,
                           state="disabled", wrap="word")
        self.log.pack(fill="x")
        self.log.tag_configure("ok",  foreground="#1D9E75")
        self.log.tag_configure("err", foreground="#E24B4A")
        self.log.tag_configure("inf", foreground=MUTED)

        # Buttons
        btn_row = tk.Frame(self, bg=BG)
        btn_row.pack(fill="x", padx=20, pady=12)

        self.gen_btn = tk.Button(btn_row, text="Generate Contracts",
                                 font=FONT_BOLD, bg=SURFACE, fg=TEXT,
                                 relief="flat", bd=0, highlightthickness=1,
                                 highlightbackground=BORDER,
                                 activebackground=SURFACE, activeforeground=TEXT,
                                 padx=14, pady=8,
                                 cursor="hand2", command=self._generate)
        self.gen_btn.pack(side="left")

        self.review_btn = tk.Button(btn_row, text="Review & Approve →",
                                    font=FONT_BOLD, bg=GREEN_LT, fg=GREEN_TX,
                                    relief="flat", bd=0, padx=14, pady=8,
                                    cursor="hand2", state="disabled",
                                    command=lambda: self.app.switch_tab(2))
        self.review_btn.pack(side="left", padx=(10, 0))

    def apply_startup_defaults(self):
        default_template = find_default_template()
        if default_template and not self.template_picker.get():
            self.template_picker.set(default_template)

    def reset_form(self):
        self.excel_picker.set("")
        self.output_picker.set("")
        self.progress["value"] = 0
        self.progress_label.configure(text="Ready")
        self.log.configure(state="normal")
        self.log.delete("1.0", "end")
        self.log.configure(state="disabled")
        self.gen_btn.configure(state="normal", text="Generate Contracts")
        self.review_btn.configure(state="disabled", text="Review & Approve →")
        self.template_picker.set("")
        default_template = find_default_template()
        if default_template:
            self.template_picker.set(default_template)

    def _excel_changed(self, path):
        self.app.excel_path = path
        fields = build_template_fields()
        if path and os.path.exists(path):
            try:
                fields = extract_excel_fields(path)
            except Exception:
                fields = build_template_fields()
        self.app.available_fields = fields

        template_tab = getattr(self.app, "template_tab", None)
        if template_tab:
            template_tab.set_available_fields(fields)

    def _template_changed(self, path):
        self.app.template_path = path
        template_tab = getattr(self.app, "template_tab", None)
        if template_tab:
            template_tab.load_template(path)

    def _edit_template(self):
        path = self.template_picker.get().strip()
        if not path or not os.path.exists(path):
            messagebox.showerror("Template missing", "Please select a valid template file first.")
            return
        self.app.template_path = path
        self.app.template_tab.load_template(path)
        self.app.switch_tab(1)

    def _log(self, msg, tag="inf"):
        self.log.configure(state="normal")
        self.log.insert("end", msg + "\n", tag)
        self.log.see("end")
        self.log.configure(state="disabled")

    def _generate(self):
        excel   = self.excel_picker.get()
        tmpl    = self.template_picker.get()
        out_dir = self.output_picker.get()

        if not excel or not os.path.exists(excel):
            messagebox.showerror("Error", "Please select a valid Excel file.")
            return
        if not tmpl or not os.path.exists(tmpl):
            messagebox.showerror("Error", "Please select a valid template file.")
            return
        if not out_dir:
            messagebox.showerror("Error", "Please select an output folder.")
            return

        out_dir = resolve_contract_output_dir(out_dir)
        os.makedirs(out_dir, exist_ok=True)
        self.output_picker.set(out_dir)
        self.app.template_path = tmpl
        self.app.output_dir    = out_dir
        self.gen_btn.configure(state="disabled")
        self.review_btn.configure(state="disabled")
        self.log.configure(state="normal"); self.log.delete("1.0", "end"); self.log.configure(state="disabled")
        self.progress["value"] = 0

        threading.Thread(target=self._run, args=(excel, tmpl, out_dir), daemon=True).start()

    def _run(self, excel, tmpl, out_dir):
        try:
            candidates, signatories = read_excel(excel)
            if not candidates:
                self.after(0, lambda: messagebox.showerror("Error", "No candidates found in Excel."))
                return

            self.app.contracts = []
            total = len(candidates)
            unmatched = []

            for i, cand in enumerate(candidates):
                name = cand.get("Full Name") or cand.get("Candidate Name") or cand.get("Name") or f"Candidate_{i+1}"
                raw_rank = cand.get("Rank") or cand.get("rank") or ""
                rank_key = normalize_rank(raw_rank)
                if not rank_key:
                    unmatched.append(f"{name}: missing Rank in Annex B")
                    continue
                if rank_key not in signatories:
                    unmatched.append(f"{name}: no Annex C signatory match for rank '{raw_rank}'")

            if unmatched:
                preview = "\n".join(unmatched[:10])
                extra = "" if len(unmatched) <= 10 else f"\n...and {len(unmatched) - 10} more"
                raise ValueError(
                    "Every Annex B row must have a matching Rank in Annex C before generation.\n\n"
                    f"{preview}{extra}"
                )

            for i, cand in enumerate(candidates):
                name = cand.get("Full Name") or cand.get("Candidate Name") or cand.get("Name") or f"Candidate_{i+1}"

                # Build replacements — merge signatory by rank
                replacements = dict(cand)
                rank = cand.get("Rank") or cand.get("rank") or ""
                replacements.update(signatories[normalize_rank(rank)])
                replacements["Issuance Date"] = current_issuance_date()

                # Write filled docx to temp location
                safe_name = re.sub(r'[\\/*?:"<>|]', "_", name)
                tmp_docx = os.path.join(out_dir, f"_tmp_{safe_name}.docx")
                fill_template(tmpl, replacements, tmp_docx)

                self.app.contracts.append({
                    "name":     name,
                    "role":     cand.get("Job Title") or cand.get("Position") or "",
                    "docx":     tmp_docx,
                    "data":     replacements,
                    "status":   "pending",
                    "review_note": "",
                    "exported_path": "",
                })

                pct = int((i + 1) / total * 100)
                msg = f"✓ {name} — ready for review"
                self.after(0, lambda m=msg, p=pct: self._tick(m, p))

            self.after(0, self._done)

        except Exception as e:
            self.after(0, lambda: (
                self._log(f"✗ Error: {e}", "err"),
                self.gen_btn.configure(state="normal"),
                messagebox.showerror("Error", str(e))
            ))

    def _tick(self, msg, pct):
        self._log(msg, "ok")
        self.progress["value"] = pct
        self.progress_label.configure(text=f"{pct}% complete")

    def _done(self):
        self.progress["value"] = 100
        total = len(self.app.contracts)
        self.progress_label.configure(text=f"{total} of {total} contracts ready for review")
        self.gen_btn.configure(state="normal")
        self.review_btn.configure(state="normal")
        self.app.review_tab.load_contracts()
        self.app.template_tab.load_template(self.app.template_path)


# ─────────────────────────────────────────────────────────────────────────────
#  Template Tab
# ─────────────────────────────────────────────────────────────────────────────

class TemplateTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self.tmpl_path = None
        self.available_fields = []
        self._build()

    def _build(self):
        tk.Label(self, text="Edit Template", font=FONT_H, bg=BG, fg=TEXT
                 ).pack(anchor="w", padx=20, pady=(18, 2))
        tk.Label(self, text="Optional: adjust the master template text, then save to return to Generate.",
                 font=FONT_SM, bg=BG, fg=MUTED).pack(anchor="w", padx=20, pady=(0, 10))

        main = tk.Frame(self, bg=BG)
        main.pack(fill="both", expand=True, padx=20, pady=(0, 10))

        # Left: text editor
        left = tk.Frame(main, bg=BG)
        left.pack(side="left", fill="both", expand=True)

        tk.Label(left, text="Template content", font=FONT_SM, bg=BG, fg=MUTED
                 ).pack(anchor="w", pady=(0, 3))

        edit_frame = tk.Frame(left, bg=SURFACE, highlightthickness=1,
                              highlightbackground=BORDER)
        edit_frame.pack(fill="both", expand=True)

        self.editor = tk.Text(edit_frame, font=("Consolas", 10), bg=SURFACE, fg=TEXT,
                              relief="flat", bd=6, wrap="word", undo=True)
        scroll = tk.Scrollbar(edit_frame, command=self.editor.yview)
        self.editor.configure(yscrollcommand=scroll.set)
        scroll.pack(side="right", fill="y")
        self.editor.pack(fill="both", expand=True)
        self.editor.tag_configure(
            "placeholder_chip",
            background=AMBER,
            foreground=AMBER_TX,
            font=("Consolas", 10, "bold"),
            borderwidth=0,
            relief="flat",
        )
        self.editor.tag_configure(
            "placeholder_name",
            background=AMBER,
            foreground=AMBER_TX,
            font=("Consolas", 10, "bold"),
        )
        self._brace_elide_supported = True
        try:
            self.editor.tag_configure("placeholder_brace", elide=True)
        except tk.TclError:
            self._brace_elide_supported = False
            self.editor.tag_configure(
                "placeholder_brace",
                background=AMBER,
                foreground="#B1863C",
                font=("Consolas", 10, "bold"),
            )

        # Highlight placeholders as you type
        self.editor.bind("<KeyRelease>", lambda e: self._highlight())

        # Right: field sidebar
        right = tk.Frame(main, bg=BG, width=150)
        right.pack(side="right", fill="y", padx=(12, 0))
        right.pack_propagate(False)

        tk.Label(right, text="INSERT FIELD", font=("Segoe UI", 8, "bold"),
                 bg=BG, fg=MUTED).pack(anchor="w", pady=(0, 6))

        self.field_frame = tk.Frame(right, bg=BG)
        self.field_frame.pack(fill="x")

        hint = tk.Label(right, text="Fields come from Excel\nheaders plus built-ins\nlike Issuance Date.",
                        font=("Segoe UI", 8), bg=BG, fg=MUTED, justify="left")
        hint.pack(anchor="w", pady=(12, 0))

        # Bottom buttons
        btn_row = tk.Frame(self, bg=BG)
        btn_row.pack(fill="x", padx=20, pady=(0, 14))

        tk.Button(btn_row, text="Back", font=FONT_SM, bg=SURFACE, fg=TEXT,
                  relief="flat", bd=0, highlightthickness=1,
                  highlightbackground=BORDER, padx=12, pady=6,
                  cursor="hand2", command=self._go_back).pack(side="left")

        tk.Button(btn_row, text="Open in Word", font=FONT_SM, bg=SURFACE, fg=TEXT,
                  relief="flat", bd=0, highlightthickness=1,
                  highlightbackground=BORDER, padx=12, pady=6,
                  cursor="hand2", command=self._open_word).pack(side="left", padx=(10, 0))

        tk.Button(btn_row, text="Save Template", font=FONT_BOLD, bg=SURFACE, fg=TEXT,
                  relief="flat", bd=0, highlightthickness=1,
                  highlightbackground=BORDER, activebackground=SURFACE,
                  activeforeground=TEXT, padx=14, pady=6,
                  cursor="hand2", command=self._save).pack(side="left", padx=(10, 0))

        self.status_lbl = tk.Label(btn_row, text="", font=FONT_SM, bg=BG, fg=GREEN_TX)
        self.status_lbl.pack(side="left", padx=(12, 0))

    def load_template(self, path):
        self.tmpl_path = path if path and os.path.exists(path) else None
        self.editor.delete("1.0", "end")
        if not self.tmpl_path:
            return
        doc = Document(self.tmpl_path)
        text = normalize_placeholder_spacing("\n".join(p.text for p in doc.paragraphs))
        self.editor.insert("1.0", text)
        self._highlight()
        self._refresh_fields()

    def set_available_fields(self, fields):
        self.available_fields = build_template_fields(fields)
        self._refresh_fields()

    def _refresh_fields(self):
        for w in self.field_frame.winfo_children():
            w.destroy()

        fields = self.available_fields or self.app.available_fields
        if not fields:
            tk.Label(self.field_frame, text="No fields available.",
                     font=("Segoe UI", 8), bg=BG, fg=MUTED, justify="left").pack(anchor="w")
            return
        for f in fields:
            self._make_field_btn(f)

    def _make_field_btn(self, name):
        btn = tk.Button(self.field_frame, text=name, font=FONT_SM,
                        bg=AMBER, fg=AMBER_TX, relief="flat", bd=0,
                        highlightthickness=1, highlightbackground=AMBER_BD,
                        padx=8, pady=4, cursor="hand2", anchor="w",
                        command=lambda n=name: self._insert_field(n))
        btn.pack(fill="x", pady=2)

    def _insert_field(self, name):
        self.editor.insert("insert", f"{{{{ {name} }}}}")
        self._highlight()

    def _highlight(self):
        for tag in ("placeholder_chip", "placeholder_name", "placeholder_brace"):
            self.editor.tag_remove(tag, "1.0", "end")

        text = self.editor.get("1.0", "end-1c")
        for m in PLACEHOLDER_RE.finditer(text):
            start = f"1.0 + {m.start()} chars"
            open_end = f"1.0 + {m.start() + 2} chars"
            inner_start = open_end
            inner_end = f"1.0 + {m.end() - 2} chars"
            close_start = inner_end
            end = f"1.0 + {m.end()} chars"

            self.editor.tag_add("placeholder_chip", start, end)
            if m.end() - m.start() > 4:
                self.editor.tag_add("placeholder_name", inner_start, inner_end)
            self.editor.tag_add("placeholder_brace", start, open_end)
            self.editor.tag_add("placeholder_brace", close_start, end)

    def _open_word(self):
        if self.tmpl_path:
            ok, err = open_path(self.tmpl_path)
            if not ok:
                messagebox.showerror("Open failed", f"Could not open the template.\n\n{err}")

    def _go_back(self):
        self.app.show_loading(
            "Returning to Generate",
            "Opening the Generate page."
        )
        self.after(10, self._finish_go_back)

    def _finish_go_back(self):
        self.app.switch_tab(0)
        self.app.update_idletasks()
        self.app.hide_loading()

    def _save(self):
        if not self.tmpl_path:
            messagebox.showwarning("No template", "Generate contracts first to load a template.")
            return
        self.app.show_loading(
            "Saving template",
            "Updating the template and returning to Generate."
        )
        self.after(10, self._finish_save)

    def _finish_save(self):
        try:
            content = self.editor.get("1.0", "end-1c")
            save_docx_plain_text(self.tmpl_path, content)
        except Exception as exc:
            self.app.switch_tab(1)
            self.app.update_idletasks()
            self.app.hide_loading()
            messagebox.showerror("Save failed", str(exc))
            return

        self.status_lbl.configure(text="✓ Saved")
        self.after(2000, lambda: self.status_lbl.configure(text=""))
        self.app.switch_tab(0)
        self.app.update_idletasks()
        self.app.hide_loading()


# ─────────────────────────────────────────────────────────────────────────────
#  Review Tab
# ─────────────────────────────────────────────────────────────────────────────

class ReviewTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._idx = 0
        self._rows = []
        self._is_editing = False
        self._build()

    def _build(self):
        tk.Label(self, text="Review & Approve", font=FONT_H, bg=BG, fg=TEXT
                 ).pack(anchor="w", padx=20, pady=(18, 2))
        tk.Label(self, text="Review each contract, edit it in the app if needed, then approve for export.",
                 font=FONT_SM, bg=BG, fg=MUTED).pack(anchor="w", padx=20, pady=(0, 10))

        main = tk.Frame(self, bg=BG)
        main.pack(fill="both", expand=True, padx=20)

        # ── Left: candidate list ──────────────────────────────────────────
        list_outer = tk.Frame(main, bg=SURFACE, highlightthickness=1,
                              highlightbackground=BORDER, width=200)
        list_outer.pack(side="left", fill="y")
        list_outer.pack_propagate(False)

        list_hdr = tk.Frame(list_outer, bg=BG)
        list_hdr.pack(fill="x")
        self.list_count = tk.Label(list_hdr, text="No contracts", font=FONT_SM,
                                   bg=BG, fg=MUTED, padx=10, pady=8)
        self.list_count.pack(anchor="w")
        tk.Frame(list_hdr, bg=BORDER, height=1).pack(fill="x")

        self.list_canvas = tk.Canvas(list_outer, bg=SURFACE, highlightthickness=0)
        list_scroll = tk.Scrollbar(list_outer, orient="vertical",
                                   command=self.list_canvas.yview)
        self.list_canvas.configure(yscrollcommand=list_scroll.set)
        list_scroll.pack(side="right", fill="y")
        self.list_canvas.pack(fill="both", expand=True)
        self.list_inner = tk.Frame(self.list_canvas, bg=SURFACE)
        self.list_window = self.list_canvas.create_window((0, 0), window=self.list_inner, anchor="nw")
        self.list_inner.bind("<Configure>",
            lambda e: self.list_canvas.configure(
                scrollregion=self.list_canvas.bbox("all")))
        self.list_canvas.bind("<Configure>", self._resize_list_inner)

        # ── Right: preview ───────────────────────────────────────────────
        right = tk.Frame(main, bg=BG)
        right.pack(side="left", fill="both", expand=True, padx=(12, 0))

        # Toolbar
        toolbar = tk.Frame(right, bg=SURFACE, highlightthickness=1,
                           highlightbackground=BORDER)
        toolbar.pack(fill="x")

        self.preview_name = tk.Label(toolbar, text="—", font=FONT_BOLD,
                                     bg=SURFACE, fg=TEXT, padx=12, pady=8)
        self.preview_name.pack(side="left")

        # Nav
        nav = tk.Frame(toolbar, bg=SURFACE)
        nav.pack(side="left", padx=8)
        tk.Button(nav, text="‹", font=FONT, bg=SURFACE, fg=TEXT, relief="flat",
                  bd=0, highlightthickness=1, highlightbackground=BORDER,
                  width=2, cursor="hand2", command=self._prev).pack(side="left")
        self.nav_label = tk.Label(nav, text="— / —", font=FONT_SM, bg=SURFACE,
                                  fg=MUTED, padx=6)
        self.nav_label.pack(side="left")
        tk.Button(nav, text="›", font=FONT, bg=SURFACE, fg=TEXT, relief="flat",
                  bd=0, highlightthickness=1, highlightbackground=BORDER,
                  width=2, cursor="hand2", command=self._next).pack(side="left")

        self.edit_btn = SolidActionButton(toolbar, text="Edit", font=FONT_SM,
                                          bg=RED_BTN, fg=RED_TX,
                                          activebackground=RED_BTN_ACTIVE,
                                          activeforeground=RED_TX,
                                          disabledforeground="#BD9895",
                                          disabledbackground="#F6ECEA",
                                          highlightbackground=RED_BD,
                                          padx=10, pady=4, cursor="hand2",
                                          command=self._edit_selected_contract)
        self.edit_btn.pack(side="right", padx=4, pady=6)

        self.approve_btn = SolidActionButton(toolbar, text="Approve", font=FONT_BOLD,
                                             bg=GREEN_BTN, fg=GREEN_TX,
                                             activebackground=GREEN_BTN_ACTIVE,
                                             activeforeground=GREEN_TX,
                                             disabledforeground="#8EA89F",
                                             disabledbackground="#EBF4F0",
                                             highlightbackground=GREEN_BD,
                                             padx=10, pady=4, cursor="hand2",
                                             command=self._approve_selected)
        self.approve_btn.pack(side="right", padx=4, pady=6)

        # Document preview
        doc_frame = tk.Frame(right, bg=SURFACE, highlightthickness=1,
                             highlightbackground=BORDER)
        doc_frame.pack(fill="both", expand=True, pady=(8, 0))

        self.doc_text = tk.Text(doc_frame, font=("Georgia", 10), bg=SURFACE, fg=TEXT,
                                relief="flat", bd=12, wrap="word", state="disabled",
                                cursor="arrow")
        doc_scroll = tk.Scrollbar(doc_frame, command=self.doc_text.yview)
        self.doc_text.configure(yscrollcommand=doc_scroll.set)
        doc_scroll.pack(side="right", fill="y")
        self.doc_text.pack(fill="both", expand=True)
        self.doc_text.tag_configure("filled", background=AMBER, foreground=TEXT,
                                    font=("Georgia", 10, "bold"))
        self.doc_text.tag_configure("heading", font=("Georgia", 10, "bold"))
        self.doc_text.tag_configure("title", font=("Georgia", 12, "bold"))

        # ── Summary bar ──────────────────────────────────────────────────
        summary = tk.Frame(self, bg=SURFACE, highlightthickness=1,
                           highlightbackground=BORDER)
        summary.pack(fill="x", padx=20, pady=(8, 14))

        self.summary_labels = {}
        self.summary_names = {}
        for status, colour, label in [
            ("approved", "#1D9E75", "approved"),
            ("needs_changes", "#E24B4A", "needs changes"),
            ("pending",  BORDER,    "pending"),
        ]:
            dot = tk.Canvas(summary, width=10, height=10, bg=SURFACE,
                            highlightthickness=0)
            dot.create_oval(1, 1, 9, 9, fill=colour, outline=colour)
            dot.pack(side="left", padx=(12, 4), pady=8)
            lbl = tk.Label(summary, text=f"0 {label}", font=FONT_SM,
                           bg=SURFACE, fg=MUTED)
            lbl.pack(side="left", padx=(0, 10))
            self.summary_labels[status] = lbl
            self.summary_names[status] = label

        self.export_btn = tk.Button(summary, text="Export approved →",
                                    font=FONT_BOLD, bg=SURFACE, fg=TEXT,
                                    relief="flat", bd=0, highlightthickness=1,
                                    highlightbackground=BORDER,
                                    activebackground=SURFACE, activeforeground=TEXT,
                                    padx=14, pady=6,
                                    cursor="hand2", state="disabled",
                                    command=self._export)
        self.export_btn.pack(side="right", padx=12, pady=6)

    # ── Data loading ──────────────────────────────────────────────────────

    def _resize_list_inner(self, event):
        self.list_canvas.itemconfigure(self.list_window, width=event.width)

    def load_contracts(self):
        for w in self.list_inner.winfo_children():
            w.destroy()
        self._rows = []

        contracts = self.app.contracts or []
        self.list_count.configure(text=f"{len(contracts)} contracts")

        for i, c in enumerate(contracts):
            row = tk.Frame(self.list_inner, bg=SURFACE, cursor="hand2")
            row.pack(fill="x")
            divider = tk.Frame(self.list_inner, bg=BORDER, height=1)
            divider.pack(fill="x")

            dot = StatusDot(row, status=c["status"])
            dot.pack(side="left", padx=(10, 6), pady=8)

            info = tk.Frame(row, bg=SURFACE)
            info.pack(side="left", fill="x", expand=True, pady=6)
            name_lbl = tk.Label(info, text=c["name"], font=FONT_BOLD, bg=SURFACE, fg=TEXT,
                                anchor="w")
            name_lbl.pack(anchor="w", fill="x")
            role_lbl = tk.Label(info, text=c["role"], font=FONT_SM, bg=SURFACE, fg=MUTED,
                                anchor="w")
            role_lbl.pack(anchor="w", fill="x")

            idx = i
            row.bind("<Button-1>", lambda e, i=idx: self._select(i))
            for child in row.winfo_children():
                child.bind("<Button-1>", lambda e, i=idx: self._select(i))
            for child in (name_lbl, role_lbl):
                child.bind("<Button-1>", lambda e, i=idx: self._select(i))

            self._rows.append({
                "frame": row,
                "dot": dot,
                "info": info,
                "divider": divider,
                "labels": [name_lbl, role_lbl],
            })

        if contracts:
            self._select(0)
        else:
            self.preview_name.configure(text="—")
            self.nav_label.configure(text="— / —")
            self._render_contract(None)
            self._update_review_context(None)
        self._update_summary()

    def reset_view(self):
        self._idx = 0
        self._is_editing = False
        self.load_contracts()

    def _select(self, idx):
        if self._is_editing and idx != self._idx:
            messagebox.showwarning("Save changes first", "Save the current contract before switching to another one.")
            return
        self._idx = idx
        contracts = self.app.contracts or []

        for i, r in enumerate(self._rows):
            row_bg = PURPLE_LT if i == idx else SURFACE
            r["frame"].configure(bg=row_bg)
            r["info"].configure(bg=row_bg)
            r["dot"].configure(bg=row_bg)
            for label in r["labels"]:
                label.configure(bg=row_bg)

        c = contracts[idx]
        self.preview_name.configure(text=c["name"])
        self.nav_label.configure(text=f"{idx+1} / {len(contracts)}")
        self._render_contract(c)
        self._update_review_context(c)

    def _build_highlight_pattern(self, contract):
        if not contract:
            return None

        values = []
        seen = set()
        for raw in (contract.get("data") or {}).values():
            value = re.sub(r"\s+", " ", str(raw or "").strip())
            if len(value) < 3:
                continue
            if not re.search(r"[A-Za-z0-9]", value):
                continue
            norm = value.casefold()
            if norm in seen:
                continue
            seen.add(norm)
            values.append(value)

        if not values:
            return None

        parts = []
        for value in sorted(values, key=lambda item: (-len(item), item.casefold())):
            escaped = re.escape(value).replace(r"\ ", r"\s+")
            prefix = r"(?<!\w)" if value[:1].isalnum() else ""
            suffix = r"(?!\w)" if value[-1:].isalnum() else ""
            parts.append(f"{prefix}{escaped}{suffix}")

        return re.compile("|".join(parts))

    def _insert_preview_text(self, text, *, base_tag="", highlight_re=None):
        if not text:
            return

        if not highlight_re:
            if base_tag:
                self.doc_text.insert("end", text, base_tag)
            else:
                self.doc_text.insert("end", text)
            return

        cursor = 0
        for match in highlight_re.finditer(text):
            start, end = match.span()
            if start > cursor:
                segment = text[cursor:start]
                if base_tag:
                    self.doc_text.insert("end", segment, base_tag)
                else:
                    self.doc_text.insert("end", segment)

            match_text = text[start:end]
            tags = (base_tag, "filled") if base_tag else ("filled",)
            self.doc_text.insert("end", match_text, tags)
            cursor = end

        if cursor < len(text):
            segment = text[cursor:]
            if base_tag:
                self.doc_text.insert("end", segment, base_tag)
            else:
                self.doc_text.insert("end", segment)

    def _render_contract(self, c):
        self.doc_text.configure(state="normal")
        self.doc_text.delete("1.0", "end")

        if not c:
            self.doc_text.insert("end", "No contract selected.")
        elif os.path.exists(c["docx"]):
            doc = Document(c["docx"])
            highlight_re = self._build_highlight_pattern(c)
            for para in doc.paragraphs:
                text = para.text
                if not text.strip():
                    self.doc_text.insert("end", "\n")
                    continue
                # Detect headings (ALL CAPS or numbered like "1. TITLE")
                is_heading = re.match(r"^\d+\.\s+[A-Z\s]+$", text.strip()) or text.isupper()
                tag = "heading" if is_heading else ""
                if is_heading:
                    self.doc_text.insert("end", "\n")

                self._insert_preview_text(text, base_tag=tag, highlight_re=highlight_re)
                self.doc_text.insert("end", "\n")
        elif c.get("exported_path"):
            self.doc_text.insert(
                "end",
                "This approved contract has already been exported.\n\n"
                f"Exported file:\n{c['exported_path']}\n\n"
                "The temporary review copy was cleaned up after export.",
            )
        else:
            self.doc_text.insert("end", "Generated contract file not found.")

        self.doc_text.configure(state="disabled")

    # ── Actions ──────────────────────────────────────────────────────────

    def _set_status(self, status, *, advance=False, review_note=None, clear_note=False):
        if not self.app.contracts:
            return
        contract = self.app.contracts[self._idx]
        contract["status"] = status
        if review_note is not None:
            contract["review_note"] = review_note
        elif clear_note:
            contract["review_note"] = ""

        self._rows[self._idx]["dot"].set(status)
        self._update_summary()
        self._update_review_context(contract)
        if advance:
            contracts = self.app.contracts
            for i in range(self._idx + 1, len(contracts)):
                if contracts[i]["status"] == "pending":
                    self._select(i)
                    return

    def _edit_selected_contract(self):
        if not self.app.contracts:
            return
        contract = self.app.contracts[self._idx]
        path = contract.get("docx", "")
        if not path or not os.path.exists(path):
            messagebox.showerror("Missing file", "The generated contract file could not be found.")
            return

        if self._is_editing:
            self._save_current_edit()
            return

        self._set_status("needs_changes")
        self._is_editing = True
        self.doc_text.configure(state="normal", cursor="xterm")
        self.doc_text.delete("1.0", "end")
        self.doc_text.insert("1.0", read_docx_plain_text(path))
        self.edit_btn.configure(text="Save Changes", bg=GREEN_BTN, fg=GREEN_TX,
                                activebackground=GREEN_BTN_ACTIVE,
                                activeforeground=GREEN_TX,
                                highlightbackground=GREEN_BD)
        self.approve_btn.configure(state="disabled")
        self.doc_text.focus_set()

    def _save_current_edit(self):
        if not self.app.contracts:
            return False
        contract = self.app.contracts[self._idx]
        path = contract.get("docx", "")
        if not path or not os.path.exists(path):
            messagebox.showerror("Missing file", "The generated contract file could not be found.")
            return False

        try:
            save_docx_plain_text(path, self.doc_text.get("1.0", "end-1c"))
        except Exception as exc:
            messagebox.showerror("Save failed", str(exc))
            return False

        self._is_editing = False
        self._render_contract(contract)
        self._update_review_context(contract)
        return True

    def _approve_selected(self):
        if self._is_editing and not self._save_current_edit():
            return
        self._set_status("approved", advance=True, clear_note=True)

    def _prev(self):
        if self._is_editing:
            messagebox.showwarning("Save changes first", "Save the current contract before switching to another one.")
            return
        if self._idx > 0:
            self._select(self._idx - 1)

    def _next(self):
        if self._is_editing:
            messagebox.showwarning("Save changes first", "Save the current contract before switching to another one.")
            return
        if self.app.contracts and self._idx < len(self.app.contracts) - 1:
            self._select(self._idx + 1)

    def _update_summary(self):
        contracts = self.app.contracts or []
        counts = {"approved": 0, "needs_changes": 0, "pending": 0}
        for c in contracts:
            s = c["status"]
            counts[s] = counts.get(s, 0) + 1

        for status, lbl in self.summary_labels.items():
            lbl.configure(text=f"{counts.get(status, 0)} {self.summary_names.get(status, status)}")

        approved = counts.get("approved", 0)
        self.export_btn.configure(state="normal" if approved > 0 else "disabled")

    def _update_review_context(self, contract):
        if self._is_editing and contract:
            self.edit_btn.configure(text="Save Changes", bg=GREEN_BTN, fg=GREEN_TX,
                                    activebackground=GREEN_BTN_ACTIVE,
                                    activeforeground=GREEN_TX,
                                    highlightbackground=GREEN_BD,
                                    state="normal")
            self.approve_btn.configure(bg=GREEN_BTN, fg=GREEN_TX,
                                       activebackground=GREEN_BTN_ACTIVE,
                                       activeforeground=GREEN_TX,
                                       highlightbackground=GREEN_BD,
                                       state="disabled")
            self.doc_text.configure(cursor="xterm")
            return

        has_contract = contract is not None
        has_editable_doc = bool(contract and contract.get("docx") and os.path.exists(contract["docx"]))
        self.edit_btn.configure(text="Edit", bg=RED_BTN, fg=RED_TX,
                                activebackground=RED_BTN_ACTIVE,
                                activeforeground=RED_TX,
                                highlightbackground=RED_BD,
                                state="normal" if has_editable_doc else "disabled")
        self.approve_btn.configure(bg=GREEN_BTN, fg=GREEN_TX,
                                   activebackground=GREEN_BTN_ACTIVE,
                                   activeforeground=GREEN_TX,
                                   highlightbackground=GREEN_BD,
                                   state="normal" if has_contract else "disabled")
        self.doc_text.configure(cursor="arrow")

    # ── Export ────────────────────────────────────────────────────────────

    def _export(self):
        out_dir = self.app.output_dir
        if not out_dir:
            chosen_dir = filedialog.askdirectory(title="Select export folder")
            if not chosen_dir:
                return
            out_dir = resolve_contract_output_dir(chosen_dir)
            os.makedirs(out_dir, exist_ok=True)
            self.app.output_dir = out_dir
            generate_tab = getattr(self.app, "generate_tab", None)
            if generate_tab:
                generate_tab.output_picker.set(out_dir)

        approved = [c for c in (self.app.contracts or []) if c["status"] == "approved"]
        if not approved:
            messagebox.showinfo("Nothing to export", "No contracts have been approved.")
            return

        self.export_btn.configure(state="disabled", text="Exporting…")
        threading.Thread(target=self._do_export, args=(approved, out_dir), daemon=True).start()

    def _do_export(self, approved, out_dir):
        errors = []
        exported = []

        for c in approved:
            try:
                safe = re.sub(r'[\\/*?:"<>|]', "_", c["name"])
                source_docx = c.get("docx", "")
                exported_path = c.get("exported_path", "")

                if source_docx and not os.path.exists(source_docx):
                    source_docx = ""

                if not source_docx and exported_path and os.path.exists(exported_path):
                    if exported_path.lower().endswith(".pdf"):
                        final_pdf = os.path.join(out_dir, f"Contract_{safe}.pdf")
                        if os.path.abspath(exported_path) != os.path.abspath(final_pdf):
                            shutil.copy2(exported_path, final_pdf)
                        exported.append(final_pdf)
                        continue
                    if exported_path.lower().endswith(".docx"):
                        source_docx = exported_path

                if not source_docx:
                    raise FileNotFoundError("No exported or temporary contract file is available for export.")

                final_docx = os.path.join(out_dir, f"Contract_{safe}.docx")
                if os.path.abspath(source_docx) != os.path.abspath(final_docx):
                    shutil.copy2(source_docx, final_docx)

                # Try PDF conversion
                try:
                    pdf_path = docx_to_pdf(final_docx, out_dir)
                    if os.path.exists(final_docx):
                        os.remove(final_docx)
                    c["exported_path"] = pdf_path
                    if source_docx and os.path.basename(source_docx).startswith("_tmp_") and os.path.exists(source_docx):
                        try:
                            os.remove(source_docx)
                        except Exception as cleanup_err:
                            errors.append(f"{c['name']}: exported but could not delete temporary file ({cleanup_err})")
                    if source_docx == c.get("docx"):
                        c["docx"] = ""
                    exported.append(pdf_path)
                except Exception:
                    # LibreOffice not available — keep docx
                    c["exported_path"] = final_docx
                    c["docx"] = final_docx
                    if source_docx and os.path.basename(source_docx).startswith("_tmp_") and os.path.exists(source_docx):
                        try:
                            os.remove(source_docx)
                        except Exception as cleanup_err:
                            errors.append(f"{c['name']}: exported but could not delete temporary file ({cleanup_err})")
                    exported.append(final_docx)

            except Exception as e:
                errors.append(f"{c['name']}: {e}")

        self.after(0, lambda: self._export_done(exported, errors, out_dir))

    def _export_done(self, exported, errors, out_dir):
        self.export_btn.configure(state="normal", text="Export approved →")
        session_reset = bool(exported) and not errors
        if self.app.contracts and 0 <= self._idx < len(self.app.contracts):
            current = self.app.contracts[self._idx]
            self._render_contract(current)
            self._update_review_context(current)
        self.app.show_export_result(exported, errors, out_dir, session_reset=session_reset)


class ExportResultTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._session_reset = False
        self._build()

    def _build(self):
        tk.Label(self, text="Export Result", font=FONT_H, bg=BG, fg=TEXT
                 ).pack(anchor="w", padx=20, pady=(18, 2))
        tk.Label(self, text="Your approved contracts have been prepared for delivery.",
                 font=FONT_SM, bg=BG, fg=MUTED).pack(anchor="w", padx=20, pady=(0, 12))

        card = tk.Frame(self, bg=SURFACE, highlightthickness=1,
                        highlightbackground=BORDER)
        card.pack(fill="x", padx=20, pady=(0, 10))

        self.result_title = tk.Label(card, text="Export successful", font=("Segoe UI", 14, "bold"),
                                     bg=SURFACE, fg=TEXT, anchor="w")
        self.result_title.pack(fill="x", padx=18, pady=(18, 6))

        self.result_summary = tk.Label(card, text="", font=FONT, bg=SURFACE, fg=TEXT,
                                       justify="left", anchor="w", wraplength=760)
        self.result_summary.pack(fill="x", padx=18, pady=(0, 10))

        self.result_errors = tk.Label(card, text="", font=FONT_SM, bg=SURFACE, fg=RED_TX,
                                      justify="left", anchor="w", wraplength=760)
        self.result_errors.pack(fill="x", padx=18, pady=(0, 14))

        btn_row = tk.Frame(self, bg=BG)
        btn_row.pack(fill="x", padx=20, pady=(0, 14))

        self.open_folder_btn = tk.Button(btn_row, text="Open Folder", font=FONT_BOLD,
                                         bg=SURFACE, fg=TEXT, relief="flat", bd=0,
                                         highlightthickness=1, highlightbackground=BORDER,
                                         activebackground=SURFACE, activeforeground=TEXT,
                                         padx=14, pady=8, cursor="hand2",
                                         command=self._open_folder)
        self.open_folder_btn.pack(side="left")

        self.next_btn = tk.Button(btn_row, text="Start New Batch", font=FONT_SM,
                                  bg=SURFACE, fg=TEXT, relief="flat", bd=0,
                                  highlightthickness=1, highlightbackground=BORDER,
                                  activebackground=SURFACE, activeforeground=TEXT,
                                  padx=14, pady=8, cursor="hand2",
                                  command=self._go_next)
        self.next_btn.pack(side="left", padx=(10, 0))

    def set_result(self, exported, errors, out_dir, *, session_reset=False):
        count = len(exported)
        self.app.last_export_dir = out_dir
        self._session_reset = session_reset
        self.result_title.configure(
            text="Export successful" if not errors else "Export completed with notes"
        )
        reset_note = "\nStart a new batch when you're ready." if session_reset else ""
        self.result_summary.configure(text=f"{count} contract(s) exported.\nFolder: {out_dir}{reset_note}")
        if errors:
            preview = "\n".join(errors[:5])
            extra = "" if len(errors) <= 5 else f"\n...and {len(errors) - 5} more"
            self.result_errors.configure(text=f"Issues:\n{preview}{extra}")
            self.next_btn.configure(text="Back to Review")
        else:
            self.result_errors.configure(text="")
            self.next_btn.configure(text="Start New Batch")

    def _open_folder(self):
        out_dir = self.app.last_export_dir
        if not out_dir:
            return
        ok, err = open_path(out_dir)
        if not ok:
            messagebox.showerror("Open failed", f"Could not open the export folder.\n\n{err}")

    def _go_next(self):
        if self._session_reset:
            self.app.start_new_batch()
        else:
            self.app.switch_tab(2)


class LoadingTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._build()

    def _build(self):
        shell = tk.Frame(self, bg=BG)
        shell.pack(fill="both", expand=True, padx=20, pady=20)

        card = tk.Frame(shell, bg=SURFACE, highlightthickness=1,
                        highlightbackground=BORDER)
        card.place(relx=0.5, rely=0.42, anchor="center", width=420, height=180)

        self.title_lbl = tk.Label(card, text="Preparing new batch", font=("Segoe UI", 14, "bold"),
                                  bg=SURFACE, fg=TEXT)
        self.title_lbl.pack(pady=(26, 8))

        self.detail_lbl = tk.Label(card, text="Resetting the workspace and loading the default template.",
                                   font=FONT_SM, bg=SURFACE, fg=MUTED, justify="center", wraplength=320)
        self.detail_lbl.pack(padx=20, pady=(0, 16))

        self.progress = ttk.Progressbar(card, mode="indeterminate", length=220)
        self.progress.pack()

    def show_message(self, title, detail):
        self.title_lbl.configure(text=title)
        self.detail_lbl.configure(text=detail)
        self.progress.start(10)

    def stop(self):
        self.progress.stop()


# ─────────────────────────────────────────────────────────────────────────────
#  Main App
# ─────────────────────────────────────────────────────────────────────────────

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Contract Generator")
        self.geometry("860x640")
        self.minsize(700, 520)
        self.configure(bg=BG)

        # Shared state
        self.contracts    = []
        self.excel_path    = ""
        self.template_path = ""
        self.output_dir    = ""
        self.available_fields = build_template_fields()
        self.last_export_dir = ""

        self._build()

    def _build(self):
        # Title bar
        hdr = tk.Frame(self, bg=SURFACE, highlightthickness=1,
                       highlightbackground=BORDER)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Contract Generator", font=("Segoe UI", 11, "bold"),
                 bg=SURFACE, fg=TEXT, padx=16, pady=10).pack(side="left")

        # Workflow steps
        step_bar = tk.Frame(self, bg=BG)
        step_bar.pack(fill="x", padx=20, pady=(14, 8))

        tk.Label(step_bar, text="Workflow", font=FONT_SM, bg=BG, fg=MUTED
                 ).pack(anchor="w", pady=(0, 6))

        step_row = tk.Frame(step_bar, bg=BG)
        step_row.pack(fill="x")

        self._step_btns = []
        self._step_targets = [0, 2]
        self._tab_step_map = {0: 0, 1: 0, 2: 1, 3: 1}
        self._tabs = []
        self._active = 0

        names = ["1. Generate", "2. Review & Approve"]
        for i, name in enumerate(names):
            btn = tk.Button(step_row, text=name, font=FONT_SM,
                            bg=SURFACE, fg=MUTED, relief="flat", bd=0,
                            highlightthickness=1, highlightbackground=BORDER,
                            activebackground=SURFACE, activeforeground=TEXT,
                            padx=16, pady=9, cursor="hand2",
                            command=lambda target=self._step_targets[i]: self.switch_tab(target))
            btn.pack(side="left")
            self._step_btns.append(btn)
            if i < len(names) - 1:
                tk.Label(step_row, text=">", font=FONT_SM, bg=BG, fg=MUTED
                         ).pack(side="left", padx=8)

        # Content area
        self.content = tk.Frame(self, bg=BG)
        self.content.pack(fill="both", expand=True)

        self.generate_tab = GenerateTab(self.content, self)
        self.template_tab = TemplateTab(self.content, self)
        self.review_tab   = ReviewTab(self.content, self)
        self.export_result_tab = ExportResultTab(self.content, self)
        self.loading_tab = LoadingTab(self.content, self)
        self.loading_tab.place_forget()
        self._tabs = [self.generate_tab, self.template_tab, self.review_tab, self.export_result_tab]

        self.generate_tab.apply_startup_defaults()
        self.template_tab.set_available_fields(self.available_fields)
        self.switch_tab(0)

    def switch_tab(self, idx):
        self._active = idx
        for i, tab in enumerate(self._tabs):
            if i == idx:
                tab.pack(fill="both", expand=True)
            else:
                tab.pack_forget()

        active_step = self._tab_step_map.get(idx, 0)
        if idx == 3 and getattr(self.export_result_tab, "_session_reset", False):
            active_step = 0
        for i, btn in enumerate(self._step_btns):
            if i == active_step:
                btn.configure(bg=PURPLE_LT, fg=PURPLE,
                              font=("Segoe UI", 9, "bold"),
                              highlightbackground=PURPLE_LT,
                              activebackground=PURPLE_LT,
                              activeforeground=PURPLE)
            else:
                btn.configure(bg=SURFACE, fg=MUTED,
                              font=FONT_SM,
                              highlightbackground=BORDER,
                              activebackground=SURFACE,
                              activeforeground=TEXT)

    def show_loading(self, title, detail):
        self.loading_tab.show_message(title, detail)
        self.loading_tab.place(in_=self.content, x=0, y=0, relwidth=1, relheight=1)
        self.loading_tab.lift()
        self.update_idletasks()

    def hide_loading(self):
        self.loading_tab.stop()
        self.loading_tab.place_forget()

    def reset_for_new_batch(self):
        self.contracts = []
        self.excel_path = ""
        self.output_dir = ""
        self.available_fields = build_template_fields()
        self.template_path = ""
        self.generate_tab.reset_form()
        self.template_tab.set_available_fields(self.available_fields)
        self.review_tab.reset_view()

    def start_new_batch(self):
        self.show_loading(
            "Preparing new batch",
            "Resetting the workspace and taking you back to Generate."
        )
        self.after(80, self._finish_start_new_batch)

    def _finish_start_new_batch(self):
        self.reset_for_new_batch()
        self.export_result_tab._session_reset = False
        self.switch_tab(0)
        self.update_idletasks()
        self.hide_loading()

    def show_export_result(self, exported, errors, out_dir, *, session_reset=False):
        self.export_result_tab.set_result(exported, errors, out_dir, session_reset=session_reset)
        self.switch_tab(3)


if __name__ == "__main__":
    app = App()
    app.mainloop()
